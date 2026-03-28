# 依赖安装：
# pip install python-docx

import argparse
import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class PaperFormatter:
    def __init__(self, input_file):
        self.input_file = os.path.abspath(input_file)
        base = os.path.splitext(self.input_file)[0]
        self.output_docx = base + ".formatted.docx"
        self.output_md = base + ".md"
        self.doc = Document()
        self.title = os.path.splitext(os.path.basename(self.input_file))[0]

        # 页面样式和论文首页
        self._set_style()
        self._configure_page_numbers()

    def _set_style(self):
        """设置论文默认排版样式"""
        style = self.doc.styles['Normal']
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        style.font.size = Pt(12)
        for section in self.doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)
            section.header_distance = Cm(1.5)
            section.footer_distance = Cm(1.75)

    def _configure_page_numbers(self):
        """页脚居中添加页码"""
        section = self.doc.sections[0]
        section.start_type = WD_SECTION.NEW_PAGE
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _add_title_block(self):
        """论文首页标题区"""
        title_paragraph = self.doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format = title_paragraph.paragraph_format
        title_format.space_after = Pt(18)
        run = title_paragraph.add_run(self.title)
        run.font.size = Pt(18)
        run.bold = True
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        meta_paragraph = self.doc.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_format = meta_paragraph.paragraph_format
        meta_format.space_after = Pt(18)
        run = meta_paragraph.add_run("作者：无名")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    def _add_heading(self, text, level):
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt = paragraph.paragraph_format
        fmt.first_line_indent = Cm(0)
        if level == 1:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_before = Pt(12)
            fmt.space_after = Pt(12)
            size = Pt(16)
        elif level == 2:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.space_before = Pt(10)
            fmt.space_after = Pt(6)
            size = Pt(14)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.space_before = Pt(8)
            fmt.space_after = Pt(4)
            size = Pt(12)

        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = size
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    @staticmethod
    def _heading_level(text):
        if re.fullmatch(r"(摘要|Abstract|关键词|Key words|参考文献|致谢)", text, re.IGNORECASE):
            return 1
        if re.fullmatch(r"[一二三四五六七八九十]+、.+", text):
            return 1
        if re.fullmatch(r"\d+\s+.+", text):
            return 1
        if re.fullmatch(r"\d+\.\d+\s+.+", text):
            return 2
        if re.fullmatch(r"\d+\.\d+\.\d+\s+.+", text):
            return 3
        return 0

    def read_docx(self):
        doc = Document(self.input_file)
        all_paragraphs = []
        unique_paragraphs = []
        seen = set()

        for p in doc.paragraphs:
            text = p.text
            normalized = self._normalize_internal_spaces(text.strip())
            if not normalized:
                continue
            all_paragraphs.append(normalized)

            if normalized in seen:
                continue

            seen.add(normalized)
            unique_paragraphs.append(normalized)
        return all_paragraphs, unique_paragraphs

    @staticmethod
    def _deduplicate_paragraphs(paragraphs):
        """按首次出现顺序去重段落。"""
        seen = set()
        unique = []
        for text in paragraphs:
            if text not in seen:
                seen.add(text)
                unique.append(text)
        return unique

    @staticmethod
    def _normalize_internal_spaces(text):
        """清理段落内部多余空格，保留英文单词间必要空格。"""
        text = text.replace("\u3000", " ").replace("\xa0", " ")
        text = re.sub(r"\s+", " ", text)
        # 中文字符之间不保留空格
        text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
        # 去掉中文标点前后的多余空格
        text = re.sub(r"\s+([，。！？：；、）》】」』])", r"\1", text)
        text = re.sub(r"([（《【「『])\s+", r"\1", text)
        return text

    def add_content(self, text_list):
        self._add_title_block()
        for text in text_list:
            level = self._heading_level(text)
            if level:
                self._add_heading(text, level)
                continue

            p = self.doc.add_paragraph()
            p.add_run(text)
            p_format = p.paragraph_format
            p_format.line_spacing = 1.5
            p_format.first_line_indent = Cm(0.74)
            p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(0)

            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    @staticmethod
    def _escape_markdown_text(text):
        if not text:
            return ""
        text = text.replace("\\", "\\\\")
        text = text.replace("`", "\\`")
        text = text.replace("*", "\\*")
        text = text.replace("_", "\\_")
        text = text.replace("[", "\\[")
        text = text.replace("]", "\\]")
        text = text.replace("|", "\\|")
        return text

    def export_markdown(self):
        title = os.path.splitext(os.path.basename(self.input_file))[0]
        try:
            lines = [f"# {self._escape_markdown_text(title)}", ""]
            doc = Document(self.input_file)
            for paragraph in doc.paragraphs:
                text = self._normalize_internal_spaces(paragraph.text.strip())
                if not text:
                    continue
                lines.append(self._escape_markdown_text(text))
                lines.append("")
            content = "\n".join(lines).strip() + "\n"
            with open(self.output_md, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"[√] MD 生成成功: {self.output_md}")
        except Exception as e:
            print("[!] MD 生成失败:", e)

    def export_docx(self):
        raw_paragraphs, paragraphs = self.read_docx()
        removed_count = len(raw_paragraphs) - len(paragraphs)
        print(f"读取到 {len(raw_paragraphs)} 个段落，去重后 {len(paragraphs)} 个（删除重复 {removed_count} 个）")
        self.add_content(paragraphs)
        self.doc.save(self.output_docx)
        print(f"[√] DOCX 生成成功: {self.output_docx}")

    def run(self, export_docx=True, export_markdown=True):
        if export_markdown:
            self.export_markdown()
        if export_docx:
            self.export_docx()


def main():
    parser = argparse.ArgumentParser(description="DOCX 文档排版并按需导出论文格式 DOCX 或 Markdown")
    parser.add_argument("file", help="输入 .docx 文件")
    parser.add_argument(
        "--format",
        choices=["docx", "markdown", "both"],
        default="both",
        help="选择输出格式：docx / markdown / both（默认）"
    )
    args = parser.parse_args()
    if not os.path.exists(args.file):
        print("文件不存在")
        return
    if os.path.splitext(args.file)[1].lower() != ".docx":
        print("请输入 .docx 文件")
        return
    app = PaperFormatter(args.file)
    export_docx = args.format in ("docx", "both")
    export_markdown = args.format in ("markdown", "both")
    app.run(export_docx=export_docx, export_markdown=export_markdown)


if __name__ == "__main__":
    main()
